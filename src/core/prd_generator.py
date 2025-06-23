from docx import Document
import markdown
from docx import Document
from tools.normalize_prd import normalize_prd
import json
import re

PRD_SECTIONS = [
    ("Introduction",            "introduction"),
    ("Customer Needs / Market / Business Model", "customer_needs"),
    ("Known Customers & Requests",      "known_requests"),
    ("Quantitative Customer Data",      "customer_data"),
    ("Business Model & Pricing",        "business_model"),
    ("Expected Results",       "expected_results"),
    ("Marketing & Communication",       "marketing_plan"),
    ("Key Metrics",            "metrics"),
    ("Terms & Copy",           "terms"),
    ("Personas",               "personas"),
    ("Functional Requirements","features"),
    ("Non‑Functional Requirements","non_functional"),
    ("Design Notes",           "design_notes"),
    ("Technical Specifications","technical_specs"),
]

def clean_agent_response(response_data):
    """Clean and extract meaningful content from agent responses"""
    if isinstance(response_data, dict):
        # Check for various content keys
        if 'text' in response_data:
            content = response_data['text']
        elif 'content' in response_data:
            content = response_data['content']
        elif 'parts' in response_data:
            # Handle Google ADK response format
            parts = response_data['parts']
            if isinstance(parts, list) and len(parts) > 0:
                content = parts[0].get('text', '')
            else:
                content = str(parts)
        else:
            content = str(response_data)
    elif isinstance(response_data, str):
        content = response_data
    else:
        content = str(response_data)
    
    # Clean up any JSON-like formatting
    if content.startswith('{') and 'content' in content:
        try:
            # Try to parse as JSON
            parsed = json.loads(content)
            if isinstance(parsed, dict):
                if 'parts' in parsed and isinstance(parsed['parts'], list):
                    content = parsed['parts'][0].get('text', content)
                elif 'text' in parsed:
                    content = parsed['text']
                elif 'content' in parsed:
                    content = parsed['content']
        except json.JSONDecodeError:
            pass
    
    # Remove metadata patterns
    content = re.sub(r"'usage_metadata':.*?'traffic_type':.*?'ON_DEMAND'.*?}", "", content, flags=re.DOTALL)
    content = re.sub(r"'invocation_id':.*?'timestamp':.*?\d+\.\d+}", "", content, flags=re.DOTALL)
    content = re.sub(r"'author': 'architect'", "", content)
    content = re.sub(r"'actions': \{.*?\}", "", content, flags=re.DOTALL)
    
    # Clean up extra whitespace and formatting
    content = re.sub(r'\s+', ' ', content).strip()
    
    return content

def extract_structured_content_from_debate(debate_history):
    """Extract structured PRD content from debate history"""
    prd_content = {
        "overview": "",
        "objectives": [],
        "user_stories": [],
        "functional_requirements": [],
        "non_functional_requirements": [],
        "success_metrics": [],
        "technical_specifications": [],
        "design_notes": []
    }
    
    if not debate_history:
        return prd_content
    
    # Process each round of debate
    for round_idx, round_data in enumerate(debate_history):
        if not isinstance(round_data, dict):
            continue
            
        for agent_key, response in round_data.items():
            # Clean the response
            content = clean_agent_response(response)
            
            # Extract content based on agent type
            agent_name = agent_key.replace('_', ' ').title()
            
            if 'ux' in agent_key.lower() or 'designer' in agent_key.lower():
                # UX Designer content
                user_stories = extract_user_stories(content)
                design_notes = extract_design_requirements(content)
                prd_content["user_stories"].extend(user_stories)
                prd_content["design_notes"].append(f"UX Perspective: {content[:300]}...")
                
            elif 'business' in agent_key.lower() or 'analyst' in agent_key.lower():
                # Business Analyst content
                objectives = extract_objectives(content)
                metrics = extract_metrics(content)
                prd_content["objectives"].extend(objectives)
                prd_content["success_metrics"].extend(metrics)
                if not prd_content["overview"]:
                    prd_content["overview"] = extract_overview(content)
                    
            elif 'backend' in agent_key.lower() or 'database' in agent_key.lower():
                # Backend/Database content
                tech_specs = extract_technical_specs(content)
                non_functional = extract_non_functional_requirements(content)
                prd_content["technical_specifications"].extend(tech_specs)
                prd_content["non_functional_requirements"].extend(non_functional)
                
            elif 'frontend' in agent_key.lower():
                # Frontend content
                functional_reqs = extract_functional_requirements(content)
                prd_content["functional_requirements"].extend(functional_reqs)
    
    # Clean up and deduplicate
    for key in prd_content:
        if isinstance(prd_content[key], list):
            # Remove duplicates and empty items
            prd_content[key] = list(set([item for item in prd_content[key] if item and len(item.strip()) > 10]))
    
    return prd_content

def extract_user_stories(content):
    """Extract user stories from content"""
    stories = []
    patterns = [
        r"As a (\w+),? I (?:want|need) to (.+?) so that (.+?)(?:\.|$)",
        r"User story:? (.+?)(?:\.|$)",
        r"(?:User|Customer) (?:needs?|wants?) (?:to )?(.+?)(?:\.|$)"
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            if isinstance(match, tuple):
                if len(match) == 3:  # Full user story format
                    stories.append(f"As a {match[0]}, I want to {match[1]} so that {match[2]}")
                else:
                    stories.append(match[0])
            else:
                stories.append(match)
    
    return stories[:5]  # Limit to 5 most relevant

def extract_objectives(content):
    """Extract objectives from content"""
    objectives = []
    patterns = [
        r"(?:objective|goal|aim)s?:?\s*(.+?)(?:\.|$)",
        r"(?:we (?:need|want|should|must) to|the goal is to|objective is to)\s*(.+?)(?:\.|$)",
        r"(?:primary|main|key) (?:objective|goal|focus|priority):?\s*(.+?)(?:\.|$)"
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        objectives.extend(matches)
    
    return [obj.strip() for obj in objectives if len(obj.strip()) > 10][:3]

def extract_metrics(content):
    """Extract success metrics from content"""
    metrics = []
    patterns = [
        r"(?:metric|measure|KPI|indicator)s?:?\s*(.+?)(?:\.|$)",
        r"(?:track|measure|monitor)\s*(.+?)(?:\.|$)",
        r"(?:success|performance) (?:metric|measure|indicator):?\s*(.+?)(?:\.|$)"
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        metrics.extend(matches)
    
    return [metric.strip() for metric in metrics if len(metric.strip()) > 5][:5]

def extract_technical_specs(content):
    """Extract technical specifications"""
    specs = []
    patterns = [
        r"(?:technology|tech|database|API|architecture|framework):?\s*(.+?)(?:\.|$)",
        r"(?:using|implement|choose|recommend)\s*(.+?)(?:\.|$)",
        r"(?:technical|system) (?:requirement|specification|decision):?\s*(.+?)(?:\.|$)"
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        specs.extend(matches)
    
    return [spec.strip() for spec in specs if len(spec.strip()) > 10][:5]

def extract_functional_requirements(content):
    """Extract functional requirements"""
    requirements = []
    patterns = [
        r"(?:feature|function|capability|requirement):?\s*(.+?)(?:\.|$)",
        r"(?:should|must|will) (?:be able to|support|provide|include)\s*(.+?)(?:\.|$)",
        r"(?:functional|system) requirement:?\s*(.+?)(?:\.|$)"
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        requirements.extend(matches)
    
    return [req.strip() for req in requirements if len(req.strip()) > 10][:5]

def extract_non_functional_requirements(content):
    """Extract non-functional requirements"""
    requirements = []
    patterns = [
        r"(?:performance|scalability|security|reliability|availability):?\s*(.+?)(?:\.|$)",
        r"(?:non-functional|quality) requirement:?\s*(.+?)(?:\.|$)",
        r"(?:must be|should be) (?:fast|secure|scalable|reliable|available)\s*(.+?)(?:\.|$)"
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        requirements.extend(matches)
    
    return [req.strip() for req in requirements if len(req.strip()) > 10][:3]

def extract_design_requirements(content):
    """Extract design requirements"""
    return f"Design considerations based on UX analysis: {content[:200]}..."

def extract_overview(content):
    """Extract project overview"""
    # Get first meaningful paragraph
    sentences = [s.strip() for s in content.split('.') if len(s.strip()) > 20]
    if sentences:
        return '. '.join(sentences[:3]) + '.'
    return content[:300] + '...' if len(content) > 300 else content

def generate_prd_docx(topic, debate_history, prd_state):
    """Enhanced PRD document generation that properly processes debate history"""
    from docx import Document
    
    # Extract structured content from debate
    structured_content = extract_structured_content_from_debate(debate_history)
    
    # Merge with existing prd_state if it has meaningful content
    if isinstance(prd_state, dict) and any(prd_state.values()):
        for key in structured_content:
            if key in prd_state and prd_state[key]:
                if isinstance(prd_state[key], list):
                    structured_content[key].extend(prd_state[key])
                elif isinstance(prd_state[key], str) and prd_state[key].strip():
                    if isinstance(structured_content[key], list):
                        structured_content[key].append(prd_state[key])
                    else:
                        structured_content[key] = prd_state[key]
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Product Requirements Document', 0)
    doc.add_paragraph(f'Project: {topic}')
    
    # Add date
    from datetime import datetime
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    # Add sections
    sections = [
        ("Project Overview", structured_content.get("overview", "To be defined based on further analysis.")),
        ("Objectives", structured_content.get("objectives", [])),
        ("User Stories", structured_content.get("user_stories", [])),
        ("Functional Requirements", structured_content.get("functional_requirements", [])),
        ("Non-Functional Requirements", structured_content.get("non_functional_requirements", [])),
        ("Success Metrics", structured_content.get("success_metrics", [])),
        ("Technical Specifications", structured_content.get("technical_specifications", [])),
        ("Design Notes", structured_content.get("design_notes", []))
    ]
    
    for section_title, content in sections:
        doc.add_heading(section_title, level=1)
        
        if isinstance(content, list):
            if content:
                for item in content:
                    if item and isinstance(item, str) and len(item.strip()) > 0:
                        doc.add_paragraph(f"• {item.strip()}", style='List Bullet')
            else:
                doc.add_paragraph("To be defined based on further analysis.")
        elif isinstance(content, str) and content.strip():
            doc.add_paragraph(content.strip())
        else:
            doc.add_paragraph("To be defined based on further analysis.")
        
        doc.add_paragraph()
    
    # Add debate summary
    doc.add_heading('Debate Summary', level=1)
    if debate_history:
        doc.add_paragraph(f'This PRD was generated from {len(debate_history)} rounds of expert debate.')
        
        for round_idx, round_data in enumerate(debate_history, 1):
            doc.add_heading(f'Round {round_idx}', level=2)
            
            if isinstance(round_data, dict):
                for agent_key, response in round_data.items():
                    agent_name = agent_key.replace('_', ' ').title()
                    doc.add_heading(agent_name, level=3)
                    
                    # Clean the response
                    response_text = clean_agent_response(response)
                    if response_text and len(response_text.strip()) > 0:
                        # Limit length for readability
                        if len(response_text) > 500:
                            response_text = response_text[:500] + "..."
                        doc.add_paragraph(response_text)
                    else:
                        doc.add_paragraph("No response provided.")
            else:
                doc.add_paragraph(f"Round data: {str(round_data)[:200]}...")
    else:
        doc.add_paragraph("No debate history available.")
    
    return doc


def generate_prd_docx_from_markdown(md: str) -> Document:
    # Very basic example: split by headers and paragraphs
    doc = Document()
    html = markdown.markdown(md)
    # parse html into Word paragraphs/headings...
    # (or use python‑docx‑compose)
    # For brevity, just dump raw text:
    for line in md.split("\n"):
        doc.add_paragraph(line)
    return doc