# tools/export_prd.py

from docx import Document

def export_prd_to_docx(prd_state, filename="Product_PRD.docx"):
    doc = Document()
    doc.add_heading("Product Requirements Document", 0)

    for section, items in prd_state.items():
        doc.add_heading(section.capitalize(), level=1)
        if isinstance(items, list):
            for item in items:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph(str(items))

    doc.save(filename)
    return filename
